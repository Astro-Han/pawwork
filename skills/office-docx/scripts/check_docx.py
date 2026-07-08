#!/usr/bin/env python3
"""Pre-delivery gate for the office-docx route: open the finished .docx with
python-docx and verify it really has the structure the task demands (real
heading styles, tables, non-empty body) before the run claims success. A
clean OK means the document opens and the required structural elements exist.

Run through uv so python-docx resolves from the mirror:
  uv run python <skill>/scripts/check_docx.py <doc.docx> [--require-heading] [--require-table] [--min-paragraphs N]
"""

import argparse
import sys

try:
    from docx import Document
except ImportError:
    print(
        "FAIL: python-docx is not importable. Run this through 'uv run python ...' from the "
        "directory that holds the office-docx pyproject.toml (uv must be on PATH).",
        file=sys.stderr,
    )
    sys.exit(2)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", help="path to the built .docx")
    parser.add_argument("--min-paragraphs", type=int, default=1, help="minimum non-empty paragraph count")
    parser.add_argument("--require-heading", action="store_true", help="need at least one real Heading-styled paragraph")
    parser.add_argument("--require-table", action="store_true", help="need at least one table")
    args = parser.parse_args()

    try:
        doc = Document(args.docx)
    except Exception as error:  # noqa: BLE001 - report any open failure as a gate failure
        print(f"FAIL: cannot open {args.docx}: {error}")
        return 1

    headings = [
        p
        for p in doc.paragraphs
        if (p.style is not None and p.style.name is not None and p.style.name.startswith("Heading"))
    ]
    # Count real body paragraphs only: a heading-only doc must not satisfy --min-paragraphs.
    body_paragraphs = [
        p
        for p in doc.paragraphs
        if p.text.strip()
        and not (p.style is not None and p.style.name is not None and p.style.name.startswith("Heading"))
    ]
    tables = doc.tables

    problems: list[str] = []
    if len(body_paragraphs) < args.min_paragraphs:
        problems.append(f"expected at least {args.min_paragraphs} body paragraph(s), found {len(body_paragraphs)}")
    if args.require_heading and len(headings) < 1:
        problems.append("task needs headings but no paragraph uses a real Heading style (do not fake headings with bold body text)")
    if args.require_table and len(tables) < 1:
        problems.append("task needs a table but no table was found")

    for problem in problems:
        print(f"FAIL: {problem}")
    if problems:
        print(f"FAIL: {len(problems)} structure gap(s). Fix the generator and rebuild; do not hand-edit the zip.")
        return 1
    print(
        f"OK: {len(body_paragraphs)} body paragraph(s), {len(headings)} heading(s), "
        f"{len(tables)} table(s) — document matches the brief."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
